#!/usr/bin/env python3
"""Render the Datachain Foundation community letter to a branded .docx.

Reads the canonical Markdown letter, converts it to a Word document styled with
the DCScan brand palette (monochrome black/white + zinc grays, with semantic
accent colours), writes the result into the user's Downloads folder, and can
optionally open macOS Mail.app with the file attached and pre-addressed.

Usage:
    python3 build_community_letter_docx.py
    python3 build_community_letter_docx.py --email            # also open Mail.app draft
    python3 build_community_letter_docx.py --source PATH --out PATH

No secrets, no network calls (except an optional pip bootstrap of python-docx).
"""
from __future__ import annotations

import argparse
import os
import re
import subprocess
import sys
from pathlib import Path

# --- dependency bootstrap ----------------------------------------------------
try:
    import docx  # noqa: F401
except ImportError:
    print("python-docx not found; installing it for the current user…", file=sys.stderr)
    subprocess.check_call([sys.executable, "-m", "pip", "install", "--user", "python-docx"])
    import docx  # noqa: F401

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

# --- DCScan brand palette (from dcscan-production/index.html :root) ----------
BLACK = RGBColor(0x00, 0x00, 0x00)
GRAY_900 = RGBColor(0x18, 0x18, 0x1B)  # body text
GRAY_700 = RGBColor(0x3F, 0x3F, 0x46)
GRAY_600 = RGBColor(0x52, 0x52, 0x5B)  # secondary text
GRAY_500 = RGBColor(0x71, 0x71, 0x7A)  # muted
GRAY_200 = "E4E4E7"                     # hairline rule / table borders (hex str)
SUCCESS = RGBColor(0x16, 0xA3, 0x4A)    # green-600, readable on white
SUCCESS_BG = "DCFCE7"
WARNING = RGBColor(0xB4, 0x54, 0x09)    # amber-700, readable on white
ERROR = RGBColor(0xDC, 0x26, 0x26)      # red-600, attacker address
ERROR_BG = "FEE2E2"
INFO_BG = "DBEAFE"

FONT_BODY = "Calibri"
FONT_MONO = "Consolas"

ATTACKER = "0xa8bD83cbb72D12209DB2Ac49D4Dc3d78E7760591"


def _shade(paragraph, hex_fill: str) -> None:
    """Apply a solid background fill to a paragraph (callout box effect)."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    pPr.append(shd)


def _bottom_border(paragraph, hex_color: str = GRAY_200, size: int = 6) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), hex_color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def _run(paragraph, text: str, *, bold=False, mono=False, color=GRAY_900,
         size=11, italic=False):
    r = paragraph.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = FONT_MONO if mono else FONT_BODY
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return r


# --- inline markdown tokenizer: handles **bold** and `code` -------------------
_INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")


def _emit_inline(paragraph, text: str, *, base_color=GRAY_900, base_size=11):
    for tok in _INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            _run(paragraph, tok[2:-2], bold=True, color=base_color, size=base_size)
        elif tok.startswith("`") and tok.endswith("`"):
            code = tok[1:-1]
            is_attacker = ATTACKER.lower() in code.lower()
            _run(paragraph, code, mono=True,
                 color=ERROR if is_attacker else GRAY_700,
                 size=base_size - 0.5, bold=is_attacker)
        else:
            _run(paragraph, tok, color=base_color, size=base_size)


def build(source: Path, out: Path) -> Path:
    md = source.read_text(encoding="utf-8").splitlines()

    doc = Document()
    # base document defaults
    normal = doc.styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(11)
    normal.font.color.rgb = GRAY_900
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)

    i = 0
    n = len(md)
    while i < n:
        line = md[i].rstrip("\n")
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # H1 — document title
        if stripped.startswith("# "):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            _run(p, stripped[2:], bold=True, color=BLACK, size=22)
            i += 1
            continue

        # H2 — section heading with hairline underline
        if stripped.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(6)
            _run(p, stripped[3:], bold=True, color=BLACK, size=14)
            _bottom_border(p)
            i += 1
            continue

        # horizontal rule
        if stripped == "---":
            p = doc.add_paragraph()
            _bottom_border(p)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # metadata block right under the title (**From:** … / **Date:** …)
        if stripped.startswith("**") and (":" in stripped) and len(stripped) < 120 \
                and any(k in stripped for k in ("From:", "Date:", "A letter")):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            _emit_inline(p, stripped, base_color=GRAY_600, base_size=10.5)
            i += 1
            continue

        # numbered list item (possibly multi-line until blank)
        m = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(6)
            _emit_inline(p, m.group(2))
            i += 1
            continue

        # signature lines at the very end (— Kazé …)
        if stripped.startswith("—") or stripped.startswith("Founder,"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            _run(p, stripped, italic=True, color=GRAY_600, size=11)
            i += 1
            continue

        # default paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.15
        _emit_inline(p, stripped)

        # semantic emphasis: the standalone "That balance is zero." line
        if "balance is" in stripped.lower() and "zero" in stripped.lower():
            for r in p.runs:
                r.font.color.rgb = SUCCESS
                r.bold = True
            _shade(p, SUCCESS_BG)

        i += 1

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


def open_mail_draft(attachment: Path, to_addr: str) -> None:
    """macOS only: open Mail.app with a pre-addressed draft + attachment."""
    if sys.platform != "darwin":
        print("--email is macOS-only (uses Mail.app). Skipping.", file=sys.stderr)
        return
    subject = "Datachain Foundation — Community Letter (2026-07-02): Security As Foundational"
    body = ("Please find attached the Datachain Foundation community letter of "
            "2026-07-02 regarding the 2026-06-22 incident, the recovery, and why "
            "security is foundational to the product.\n\n— Datachain Foundation")
    script = f'''
    tell application "Mail"
        set newMsg to make new outgoing message with properties {{subject:"{subject}", content:"{body}", visible:true}}
        tell newMsg
            make new to recipient at end of to recipients with properties {{address:"{to_addr}"}}
            tell content to make new attachment with properties {{file name:(POSIX file "{attachment}")}} at after last paragraph
        end tell
        activate
    end tell
    '''
    subprocess.run(["osascript", "-e", script], check=True)
    print(f"Opened a Mail.app draft to {to_addr} with {attachment.name} attached. "
          "Review and click Send.")


def main() -> int:
    default_src = Path(__file__).resolve().parent / "COMMUNITY_LETTER_2026-07-02_SECURITY_AS_FOUNDATIONAL.md"
    default_out = Path.home() / "Downloads" / "Datachain_Foundation_Community_Letter_2026-07-02.docx"

    ap = argparse.ArgumentParser(description="Render the community letter to a branded .docx")
    ap.add_argument("--source", type=Path, default=default_src)
    ap.add_argument("--out", type=Path, default=default_out)
    ap.add_argument("--email", action="store_true",
                    help="open a macOS Mail.app draft to contact@ongune.com with the file attached")
    ap.add_argument("--to", default="contact@ongune.com")
    args = ap.parse_args()

    if not args.source.exists():
        print(f"ERROR: source not found: {args.source}", file=sys.stderr)
        return 1

    out = build(args.source, args.out)
    print(f"Wrote {out}")

    if args.email:
        open_mail_draft(out, args.to)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
